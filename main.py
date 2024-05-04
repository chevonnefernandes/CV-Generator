from docx import Document
from docx.shared import Inches
import pyttsx3


class CVGenerator:
    def __init__(self):
        self.document = Document()
        self.name = ""
        self.email = ""
        self.linkedin_url = ""

    def generate_cv(self):
        self.speak(
            "Welcome to your personal CV generator. Make sure your CV headshot is saved as headshot.png before we continue"
        )
        self.add_profile_picture()
        self.add_contact_details()
        self.add_professional_profile()
        self.add_technical_skills()
        self.add_work_experience()
        self.add_footer()
        self.speak(f"Great job {self.name}! Your CV has been generated.")
        self.document.save("cv.docx")

    def speak(self, text):
        pyttsx3.speak(text)

    def add_profile_picture(self):
        self.document.add_picture("headshot.png", width=Inches(1.0))

    def add_contact_details(self):
        self.name = input("Enter your full name: ").capitalize()
        self.speak(f"Hi {self.name}. Next, let's add your email address")
        self.email = input("Enter your email address: ").lower()
        self.speak("Great. Now enter your LinkedIn url")
        self.linkedin_url = input("Enter your LinkedIn URL: ").lower()
        self.document.add_paragraph(f"{self.name} | {self.email} | {self.linkedin_url}")

    def add_professional_profile(self):
        self.speak(
            "Let's add a professional profile to concisely summarise your skills, experiences and career goals."
        )
        self.document.add_heading("Professional Profile")
        about_me = input("Add your professional profile here: ")
        self.document.add_paragraph(about_me)

    def add_technical_skills(self):
        self.speak(
            "Let's add some technical skills to highlight specific technical abilities and competencies."
        )
        self.document.add_heading("Technical Skills")
        while True:
            skill = input("Enter technical skill: ")
            p = self.document.add_paragraph(skill)
            p.style = "List Bullet"
            add_skill = input(
                "Do you have an additional skill you'd like to add? (Yes/No): "
            )
            if add_skill.lower() != "yes":
                break

    def add_work_experience(self):
        self.speak("Now let's add in your relevant work experience.")
        self.document.add_heading("Experience")
        while True:
            p = self.document.add_paragraph()
            company = input("Company name: ")
            start_date = input('Start date (in the format "Month YYYY"): ')
            end_date = input('End date (in the format "Month YYYY"): ')
            experience_details = input(f"Describe your role at {company}: ")
            p.add_run(f"{company} ").bold = True
            p.add_run(f"{start_date} - {end_date}\n").italic = True
            p.add_run(experience_details)
            add_exp = input(
                "Do you have an additional experience you would like to add? (Yes/No): "
            )
            if add_exp.lower() != "yes":
                break

    def add_footer(self):
        section = self.document.sections[0].footer
        p = section.paragraphs[0]
        p.text = "References available on request."


if __name__ == "__main__":
    cv_generator = CVGenerator()
    cv_generator.generate_cv()
